"""Tests for text extraction from KB documents."""

from __future__ import annotations

import csv
import io

import pytest

from app.core.exceptions import IngestionError
from app.services.rag.extractors import SUPPORTED_EXTENSIONS, extract_text


class TestExtractPlainText:
    """Test TXT and MD extraction."""

    def test_extract_txt_utf8(self):
        """UTF-8 text is returned as-is."""
        content = "Bienvenue au CRI Rabat.\nGuide de l'investisseur."
        result = extract_text(content.encode("utf-8"), "guide.txt")
        assert result == content

    def test_extract_md(self):
        """Markdown files are extracted like plain text."""
        content = "# Titre\n\nContenu du document."
        result = extract_text(content.encode("utf-8"), "README.md")
        assert result == content

    def test_extract_txt_latin1_fallback(self):
        """Latin-1 encoded text falls back gracefully."""
        content = "Résumé des procédures"
        data = content.encode("latin-1")
        result = extract_text(data, "procedures.txt")
        assert "sum" in result  # Partial match — latin-1 decoded

    def test_extract_empty_txt_raises(self):
        """Empty text file raises IngestionError."""
        with pytest.raises(IngestionError, match="No text content"):
            extract_text(b"   \n  \t  ", "empty.txt")


class TestExtractCSV:
    """Test CSV extraction."""

    def test_extract_csv_basic(self):
        """CSV rows are joined with pipe separators."""
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Nom", "Secteur", "Région"])
        writer.writerow(["Projet Alpha", "Industrie", "Rabat"])
        writer.writerow(["Projet Beta", "Tourisme", "Tanger"])
        data = buf.getvalue().encode("utf-8")

        result = extract_text(data, "projets.csv")

        assert "Nom | Secteur | Région" in result
        assert "Projet Alpha | Industrie | Rabat" in result
        assert "Projet Beta | Tourisme | Tanger" in result

    def test_extract_csv_empty_raises(self):
        """Empty CSV raises IngestionError."""
        with pytest.raises(IngestionError, match="No text content"):
            extract_text(b"", "empty.csv")


class TestExtractPDF:
    """Test PDF extraction."""

    def test_extract_pdf_basic(self):
        """Extract text from a minimal PDF."""
        # Build a minimal valid PDF with text
        try:
            import pdfplumber
        except ImportError:
            pytest.skip("pdfplumber not installed")

        # Create a minimal PDF using reportlab if available, else
        # use a raw PDF with embedded text
        pdf_bytes = _make_minimal_pdf("Guide de l'investisseur au Maroc")
        result = extract_text(pdf_bytes, "guide.pdf")
        assert "investisseur" in result.lower()

    def test_extract_pdf_empty_raises(self):
        """PDF with no text raises IngestionError."""
        # A minimal valid PDF structure with no text content
        pdf_bytes = _make_minimal_pdf("")
        with pytest.raises(IngestionError, match="No text content"):
            extract_text(pdf_bytes, "empty.pdf")


class TestExtractDocx:
    """Test DOCX extraction."""

    def test_extract_docx_basic(self):
        """Extract text from a minimal DOCX."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Procédures d'investissement")
        doc.add_paragraph("Formulaire de création d'entreprise")
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        result = extract_text(data, "procedures.docx")
        assert "investissement" in result.lower()
        assert "entreprise" in result.lower()

    def test_extract_docx_empty_raises(self):
        """DOCX with no paragraphs raises IngestionError."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()

        with pytest.raises(IngestionError, match="No text content"):
            extract_text(data, "empty.docx")


class TestUnsupportedExtension:
    """Test rejection of unsupported file types."""

    @pytest.mark.parametrize("filename", ["malware.exe", "archive.zip", "image.png", "data.json"])
    def test_unsupported_extension_raises(self, filename: str):
        """Unsupported extensions raise IngestionError."""
        with pytest.raises(IngestionError, match="Unsupported file extension"):
            extract_text(b"some data", filename)

    def test_supported_extensions_constant(self):
        """Verify supported extensions set."""
        assert SUPPORTED_EXTENSIONS == {".pdf", ".docx", ".txt", ".md", ".csv"}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_minimal_pdf(text: str) -> bytes:
    """Create a minimal PDF with the given text.

    Uses a raw PDF structure to avoid extra dependencies.
    """
    # Minimal valid PDF 1.4 with a single page and text
    if text:
        stream_content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    else:
        stream_content = ""

    stream_bytes = stream_content.encode("latin-1")
    stream_length = len(stream_bytes)

    pdf = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length {stream_length} >>
stream
{stream_content}
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000000 00000 n

trailer
<< /Size 6 /Root 1 0 R >>
startxref
0
%%EOF"""

    return pdf.encode("latin-1")
